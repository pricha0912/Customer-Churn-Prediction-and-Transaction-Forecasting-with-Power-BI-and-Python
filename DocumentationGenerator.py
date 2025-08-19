import json
import pandas as pd
from docx import Document
from docx.shared import Pt
from graphviz import Digraph
import zipfile
from typing import Any, Dict, List

# --------- File Paths ---------

vpax_file = r"C:\Users\RICHA\OneDrive\Desktop\PowerBI\model.vpax"
excel_output = "PowerBI_Documentation.xlsx"
word_output = "PowerBI_Documentation.docx"
er_diagram_output = "PowerBI_ER_Diagram"  # renders PowerBI_ER_Diagram.png

# --------- Load JSON from VPAX ---------
with zipfile.ZipFile(vpax_file, "r") as z:
    print("Contents of VPAX:", z.namelist())
    with z.open("DaxModel.json") as f:
        model = json.load(f)

# --------- Build $id → object lookup for fast $ref resolution ---------
id_map: Dict[str, Dict[str, Any]] = {}

def build_id_map(obj: Any):
    """Recursively collect every dict with a $id so we can dereference $ref quickly."""
    if isinstance(obj, dict):
        if "$id" in obj:
            id_map[obj["$id"]] = obj
        for v in obj.values():
            build_id_map(v)
    elif isinstance(obj, list):
        for item in obj:
            build_id_map(item)

build_id_map(model)

def deref(obj: Any) -> Dict[str, Any]:
    """Return the resolved object if obj is a {$ref: '...'} dict; else return obj as dict or {}."""
    if isinstance(obj, dict) and "$ref" in obj:
        return id_map.get(obj["$ref"], {})
    return obj if isinstance(obj, dict) else {}

# --------- Extract tables, columns, measures ---------
tables_list = model.get("Tables", [])
relationships_list = model.get("Relationships", [])

tables_data: List[Dict[str, Any]] = []
columns_data: List[Dict[str, Any]] = []
measures_data: List[Dict[str, Any]] = []

for t in tables_list:
    t = deref(t)
    t_name = t.get("TableName") or t.get("Name") or "Unknown"
    t_hidden = t.get("IsHidden", False)
    tables_data.append({"Table Name": t_name, "Is Hidden": t_hidden})

    # Columns
    for c in t.get("Columns", []):
        c = deref(c)
        columns_data.append({
            "Table": t_name,
            "Column": c.get("ColumnName") or c.get("Name") or "Unknown",
            "Data Type": c.get("DataType", "Unknown"),
            "Is Calculated": (c.get("ColumnType") == "Calculated") or c.get("IsCalculated", False),
            "Source Column": c.get("SourceColumn", ""),
            "Is Hidden": c.get("IsHidden", False),
        })

    # Measures
    for m in t.get("Measures", []):
        m = deref(m)
        measures_data.append({
            "Table": t_name,
            "Measure": m.get("MeasureName") or m.get("Name") or "Unknown",
            "DAX Formula": m.get("MeasureExpression") or m.get("Expression") or "",
            "Format String": m.get("FormatString", ""),
        })

# --------- Extract relationships (resolve From/To Column → Table/Column names) ---------
relationships_data: List[Dict[str, Any]] = []
for r in relationships_list:
    r = deref(r)
    from_col = deref(r.get("FromColumn"))
    to_col = deref(r.get("ToColumn"))
    from_table = deref(from_col.get("Table")) if from_col else {}
    to_table = deref(to_col.get("Table")) if to_col else {}

    # Prefer TOM-style fields (From/ToCardinalityType) when present, fall back to generic "Cardinality"
    card = (f"{r.get('FromCardinalityType','?')}-{r.get('ToCardinalityType','?')}"
            if (r.get('FromCardinalityType') or r.get('ToCardinalityType')) else r.get("Cardinality", "Unknown"))

    relationships_data.append({
        "From Table": from_table.get("TableName") or from_table.get("Name") or "Unknown",
        "From Column": from_col.get("ColumnName") or from_col.get("Name") or "Unknown",
        "To Table": to_table.get("TableName") or to_table.get("Name") or "Unknown",
        "To Column": to_col.get("ColumnName") or to_col.get("Name") or "Unknown",
        "Cardinality": card,
        "Cross Filtering": r.get("CrossFilteringBehavior", r.get("CrossFiltering", "single")),
    })

# --------- Save Excel ---------
df_tables = pd.DataFrame(tables_data)
df_columns = pd.DataFrame(columns_data)
df_measures = pd.DataFrame(measures_data)
df_relationships = pd.DataFrame(relationships_data)

with pd.ExcelWriter(excel_output, engine="openpyxl") as writer:
    df_tables.to_excel(writer, index=False, sheet_name="Tables")
    df_columns.to_excel(writer, index=False, sheet_name="Columns")
    df_measures.to_excel(writer, index=False, sheet_name="Measures")
    df_relationships.to_excel(writer, index=False, sheet_name="Relationships")
print(f"✅ Excel Documentation saved as {excel_output}")

# --------- Save Word ---------
doc = Document()
doc.add_heading("Power BI Data Model Documentation", level=0)

doc.add_heading("Tables", level=1)
for t in tables_data:
    doc.add_paragraph(f"📂 {t['Table Name']} (Hidden: {t['Is Hidden']})", style="List Bullet")

doc.add_heading("Columns", level=1)
for c in columns_data:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"{c['Table']}.{c['Column']}").bold = True
    p.add_run(f" — Type: {c['Data Type']}, Calculated: {c['Is Calculated']}, "
              f"Source: {c['Source Column']}, Hidden: {c['Is Hidden']}")

doc.add_heading("Measures", level=1)
for m in measures_data:
    doc.add_heading(f"{m['Table']} → {m['Measure']}", level=2)
    dax = doc.add_paragraph(m["DAX Formula"])
    try:
        dax.style.font.name = "Consolas"
        dax.style.font.size = Pt(9)
    except Exception:
        pass
    if m.get("Format String"):
        doc.add_paragraph(f"Format: {m['Format String']}")

doc.add_heading("Relationships", level=1)
for r in relationships_data:
    doc.add_paragraph(
        f"{r['From Table']}.{r['From Column']}  ➝  {r['To Table']}.{r['To Column']} "
        f"(Cardinality: {r['Cardinality']}, Cross Filter: {r['Cross Filtering']})",
        style="List Bullet"
    )

doc.save(word_output)
print(f"✅ Word Documentation saved as {word_output}")

# --------- ER Diagram (Graphviz) ---------
try:
    dot = Digraph(comment="Power BI ER Diagram", format="png")
    dot.attr("node", shape="box", style="rounded,filled", color="lightblue2", fontname="Helvetica")

    # Nodes: tables
    for t in tables_data:
        dot.node(t["Table Name"], t["Table Name"])

    # Edges: relationships
    for r in relationships_data:
        dot.edge(r["From Table"], r["To Table"],
                 label=f"{r['From Column']} → {r['To Column']} ({r['Cardinality']})")

    dot.render(er_diagram_output, cleanup=True)
    print(f"✅ ER Diagram saved as {er_diagram_output}.png")

except Exception as e:
    print(f"⚠️ Could not generate ER diagram: {e}")
